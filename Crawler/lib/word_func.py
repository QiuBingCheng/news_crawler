# -*- coding: utf-8 -*-
"""
Created on Tue Dec 11 09:43:24 2018

@author: User
"""

# from sklearn.feature_extraction.text import CountVectorizer
from nltk.tokenize import word_tokenize
import jieba
import configparser
from wordcloud import WordCloud
from docx import Document
import datetime
# %%
# init

config = configparser.ConfigParser()
config.read('config.ini')

jieba.set_dictionary(config['jieba']['dictionary'])

# %%
# word processing


def filter_sentecnce(content: str):
    '''回傳文本切詞結果，詞之間以空白隔開，type為String'''

    # read unwanted words
    with open(config['jieba']['stopwords'], encoding='UTF-8') as f:
        stopwords = f.readlines()
        stopwords = [w.replace('\n', " ")for w in stopwords]
        stopwords = [w.replace(" ", "") for w in stopwords]
        stopwords.append('\n')
        stopwords.append('\n   \n')
        stopwords.append('\x0b')

    # tokenize words and remove unwanted words.
    word_tokens = jieba.cut(content)
    filtered_word_tokens = [w for w in word_tokens if w not in stopwords]
    return " ".join(filtered_word_tokens)

# %%


def process_to_image(news_info: dict, filename: str):
    content_text = " ".join([news["content"] for news in news_info])
    cut_text = filter_sentecnce(content_text)  # 文本切詞結果
    wordcloud = WordCloud(collocations=False,
                          # 字體設定(是中文一定要設定，否則會是亂碼)
                          font_path=config["word_cloud"]["font_path"],
                          width=800,  # 圖片寬度
                          height=600,  # 圖片高度
                          background_color='white',  # 圖片底色
                          margin=2  # 文字之間的間距
                          ).generate(cut_text)  # 要放入的文字)

    image = wordcloud.to_image()
    image.save(f"{config['result']['image']}\\{filename}.png")


def process_to_word(news_info: dict, title: str, filename: str):

    document = Document()
    document.add_heading(title)
    p = document.add_paragraph('這是一篇由爬蟲抓取GOOGLE新聞，並以python自動寫入的WORD檔')
    p.add_run().bold = True

    for news in news_info:
        document.add_heading(news["title"], 3)
        document.add_paragraph(news["souce"])
        document.add_paragraph(news["content"])

    document.save(f"{config['storage']['doc']}\\{filename}.docx")
