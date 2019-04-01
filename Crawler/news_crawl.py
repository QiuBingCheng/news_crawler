# -*- coding: utf-8 -*-
"""
Created on Tue Dec 11 08:22:40 2018

@author: User
"""
import os
import re
path = r'C:\Users\User\Crawler'
os.chdir(path)
from WordFrequencyFunc import fileter_sentecnce,text_freq
import requests
from bs4 import BeautifulSoup
import time
from docx import Document
from wordcloud import WordCloud
import warnings
warnings.filterwarnings('ignore')
###############################################################################
#                     請使用者輸入關鍵字，並爬下的相關新聞 (Title,Source,Content)                           #
###############################################################################
def get_soup(url):
    list_req = requests.get(url)
    if list_req.status_code == requests.codes.ok:
        soup = BeautifulSoup(list_req.text, 'html.parser') # 以 BeautifulSoup 解析 HTML 程式碼
        return soup
    else:
        print("此網頁找不到!")
        return None

#不同的報社有不同的網頁排版
def get_content(url,source):
    list_req = requests.get(url)
    if list_req.status_code == requests.codes.ok:
        soup = BeautifulSoup(list_req.text, 'html.parser') # 以 BeautifulSoup 解析 HTML 程式碼
        content = []
        if source.find('Yahoo奇摩股市') != -1:
            table_tag = soup.find('table',id = 'aritcletable')
            ps_tag = table_tag.find_all('p')
      
        elif source.find('中時電子報') != -1:
            div_tag = soup.find('div',class_ = "article-body")
            ps_tag = div_tag.find_all('p')
         
        elif source.find('蘋果日報') != -1:
            div_tag = soup.find('div',class_="ndArticle_margin")
            ps_tag = div_tag.find_all('p')
          
        elif source.find('經濟日報') != -1:
            div_tag = soup.find('div',id="story_body_content")
            ps_tag = div_tag.find_all('p')
            
        else:
             ps_tag = soup.find_all('p')
             
        for p in ps_tag:
                content.append(p.text)
        content = " ".join(content)
        return content
    else:
        print("找不到此網址 :\n{}".format(url))
        return None


keyword = input('請輸入關鍵字 :')
pages = int(input("請輸入爬取的頁數 :"))
print("=======開始進行爬取動作=======")

path = path+"\\"+keyword
if not os.path.isdir(path):
    os.mkdir(path) 
os.chdir(path)

title_list = []
content_list = []
source_list = []

for i in range(pages):
    url='https://www.google.com/search?q={}&tbm=nws&ei=NGqfXOi7CMaEr7wP84yJkAg\
        &start=20&sa=N&ved=0ahUKEwjopajS96nhAhVGwosBHXNGAoIQ8tMDCFI&biw=537&bih=462&dpr=1.88'.format(keyword,i*10)
    soup = get_soup(url)   
    print("正在爬取第{}頁新聞..".format(i+1))
  # 以 CSS 的 class 抓出各類頭條新聞
    h3s = soup.find_all('h3')
    source = soup.find_all('div',class_='slp')
    
    for i in range(len(h3s)):
        content_link = h3s[i].find('a').get('href')
        content_link = content_link[content_link.find('http'):content_link.find('&')]
        content = get_content(content_link,source[i].text)
        if content:
            content_list.append(content) 
            title_list.append(h3s[i].text)
            source_list.append(source[i].text)
        time.sleep(2)
        
print("共爬取 {} 筆紀錄  ".format(len(content_list)))      
###############################################################################
#                    將新聞文本記錄在word檔中，並製作成文字雲。                  #
###############################################################################
document = Document()
document.add_heading(keyword+' NEWS', 0)
p=document.add_paragraph('這是一篇由爬蟲抓取GOOGLE新聞，並以python自動寫入的WORD檔')
p.add_run().bold=True

for i in range(len(content_list)):
    document.add_heading(title_list[i],3)
    document.add_paragraph(source_list[i])
    document.add_paragraph(content_list[i])
            

document.save(keyword+" NEWS.docx")  
print("文字檔 {} 已經儲存 ".format(keyword+" NEWS.docx"))

# =============================================================================
#                   將新聞文本製作成文字雲
# =============================================================================
#取得以jieba模組切詞後的結果  
content_text =" ".join(content_list)
cut_text= fileter_sentecnce(content_text)
#製作文字雲
wordcloud=WordCloud(collocations=False, 
                     font_path='C:\Windows\Fonts\msjhbd.ttc', # 字體設定(是中文一定要設定，否則會是亂碼)
                     #font_path='NotoSansCJKjp-Black.otf',  字體設定(是中文一定要設定，否則會是亂碼)
                     width=800, # 圖片寬度
                     height=600,  # 圖片高度
                     background_color = 'white', #圖片底色
                     margin=2 # 文字之間的間距
                     ).generate(cut_text) # 要放入的文字)
image = wordcloud.to_image()
image.show()

# 儲存圖片
filepath = keyword+' NEWS' +  '.png'
image.save(filepath,"PNG")
print("文字雲 {} 已經儲存 ".format(filepath))




