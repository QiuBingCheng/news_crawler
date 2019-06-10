# -*- coding: utf-8 -*-
"""
Created on Tue Apr 16 14:46:44 2019

@author: User
"""

from WordFrequencyFunc import fileter_sentecnce,text_freq
import pandas as pd
from qtpandas.models.DataFrameModel import DataFrameModel
import requests
from bs4 import BeautifulSoup
import time
from docx import Document
from wordcloud import WordCloud
import urllib.parse
import os
import warnings
import urllib.parse
import datetime
warnings.filterwarnings('ignore')
###############################################################################
#                     請使用者輸入關鍵字，並爬下的相關新聞 (Title,Source,Content)                           #
###############################################################################
class Crawler():
    def __init__(self,keyword,pages,thread):
        self.__keyword = keyword
        self.__pages = pages
        #self.__output_area = output_area
        self.thread = thread
        self.fig_name = ''
        self.title_list = []
        self.content_list = []
        self.source_list = []
        self.link_list = []

    def get_keyword(self):
        return self.__keyword

    def get_pages(self):
        return self.__pages

    def get_soup(self,url):
        list_req = requests.get(url)
        if list_req.status_code == requests.codes.ok:
            soup = BeautifulSoup(list_req.text, 'html.parser') # 以 BeautifulSoup 解析 HTML 程式碼
            return soup
        else:
            self.thread.update_msg_signal.emit('網頁失效...')
            return None
    #不同的報社有不同的網頁排版
    def get_content(self,url,source):
        try :   
            list_req = requests.get(url, headers={'Accept-Encoding': ''},timeout=6, allow_redirects=True)
        except:
            self.thread.update_msg_signal.emit("ContentDecodingError! \n 失效來源：{}\n 失效網址 :{}".format(source,url))
            return None 
        if list_req.status_code == requests.codes.ok:
            soup = BeautifulSoup(list_req.text, 'html.parser')
            flag = True 
            if soup.find_all('p') == [] :
                flag = False
           
            elif source.find('Yahoo奇摩股市') != -1:
                table_tag = soup.find('table',id = 'aritcletable')
                if table_tag == None :
                    flag = False
                else:
                    ps_tag = table_tag.find_all('p')

            elif source.find('中時電子報') != -1:
                div_tag = soup.find('div',class_ = "article-body")
                if div_tag == None:
                    flag = False
                else:
                    ps_tag = div_tag.find_all('p')
             
            elif source.find('蘋果日報') != -1:
                div_tag = soup.find('div',class_="ndArticle_margin")
                if div_tag == None:
                    flag = False
                else:
                    ps_tag = div_tag.find_all('p')
            
            elif source.find('香港經濟日報') != -1:
                div_tag = soup.find('div',class_='article-detail-content-container')
                if div_tag == None:
                    flag = False
                else:
                    ps_tag = div_tag.find_all('p')

            elif source.find('經濟日報') != -1:
                div_tag = soup.find('div',id="story_body_content")
                if div_tag == None:
                    flag = False
                else:
                    ps_tag = div_tag.find_all('p')

            elif source.find('DIGITIMES') != -1:
                ps_tag = soup.find_all('p',class_="main_p")
                if ps_tag == None:
                    flag = False  
            else:
                 ps_tag = soup.find_all('p')

            if flag == False:
                self.thread.update_msg_signal.emit("失效來源：{}\n 失效網址 :{}".format(source,url))
                return None

            content = []     
            for p in ps_tag:
                    content.append(p.text)
            content = " ".join(content)
            return content

        self.thread.update_msg_signal.emit("失效來源：{}\n 失效網址 :{}".format(source,url))
        return None 

    def crawl(self):
        for i in range(self.__pages):  
            url='https://www.google.com/search?q={}&tbm=nws&ei=NGqfXOi7CMaEr7wP84yJkAg\
                &start={}&sa=N&ved=0ahUKEwjopajS96nhAhVGwosBHXNGAoIQ8tMDCFI&biw=537&bih=462&dpr=1.88'.format(urllib.parse.quote(self.__keyword),i*10)
            soup = self.get_soup(url)   
          # 以 CSS 的 class 抓出各類頭條新聞
            h3s = soup.find_all('h3')
            source = soup.find_all('div',class_='slp')
            
            for i in range(len(h3s)):
                content_link = h3s[i].find('a').get('href') #取得每則新聞連結
                content_link = content_link[content_link.find('http'):content_link.find('&')]
                content_link = urllib.parse.unquote(content_link)
                self.link_list.append(content_link)
                content = self.get_content(content_link,source[i].text)
                if content:
                    self.content_list.append(content) 
                    self.title_list.append(h3s[i].text)
                    self.thread.update_msg_signal.emit(h3s[i].text)
                    self.source_list.append(source[i].text)
                time.sleep(5)
            return   

    def record_in_word(self):
        path = os.getcwd()+"\\"+self.__keyword
        if not os.path.isdir(path):
            os.mkdir(path) 
        os.chdir(path)

        document = Document()
        document.add_heading(self.__keyword+' NEWS', 0)
        p = document.add_paragraph('這是一篇由爬蟲抓取GOOGLE新聞，並以python自動寫入的WORD檔')
        p.add_run().bold=True
        
        for i in range(len(self.content_list)):
            document.add_heading(self.title_list[i],3)
            document.add_paragraph(self.source_list[i])
            document.add_paragraph(self.content_list[i])
        document.save(self.__keyword+" NEWS.docx") 

    def record_in_excel(self):
        content_text = " ".join(self.content_list)
        cut_text = fileter_sentecnce(content_text)  #文本切詞結果
        words_counts = text_freq(cut_text)
        words_counts.to_csv(self.keyword+" NEWS.csv")
 
    def content_to_wordcolud(self):
        cut_text = " ".join(self.content_list)
        cut_text = fileter_sentecnce(content_text)  #文本切詞結果
        #製作文字雲
        wordcloud=WordCloud(collocations=False, 
                             font_path='C:\Windows\Fonts\msjhbd.ttc', # 字體設定(是中文一定要設定，否則會是亂碼)
                             #font_path='NotoSansCJKjp-Black.otf',  字體設定(是中文一定要設定，否則會是亂碼)
                             width=800, # 圖片寬度
                             height=600,  # 圖片高度
                             background_color = 'white', #圖片底色
                             margin=2 # 文字之間的間距
                             ).generate(cut_text) # 要放入的文字)
        image = wordcloud.to_image()
        image.show()

        # 儲存圖片
        filepath = self.keyword+' NEWS' +  '.png'
        image.save(filepath,"PNG")
        #print("文字雲 {} 已經儲存 ".format(filepath))

    def content_to_wordcolud2(self):
        cut_text = " ".join(self.content_list)
        cut_text = fileter_sentecnce(cut_text)  #文本切詞結果
        #製作文字雲
        wordcloud=WordCloud(collocations=False, 
                             font_path='C:\Windows\Fonts\msjhbd.ttc', # 字體設定(是中文一定要設定，否則會是亂碼)
                             #font_path='NotoSansCJKjp-Black.otf',  字體設定(是中文一定要設定，否則會是亂碼)
                             width=800, # 圖片寬度
                             height=600,  # 圖片高度
                             background_color = 'white', #圖片底色
                             margin=2 # 文字之間的間距
                             ).generate(cut_text) # 要放入的文字)
        image = wordcloud.to_image()
        now = datetime.datetime.now().strftime("%Y-%m-%d %H-%M-%S")
        # 儲存圖片 檔名:2019-05-27 20-45-33_周杰倫
        self.fig_name = str(now)+"_"+self.__keyword +'.png'
        image.save("D:\\user\\PyQt5\\Crawler\\news_WordCloud\\"+self.fig_name,"PNG")
      
 
 
